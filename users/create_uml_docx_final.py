#!/usr/bin/env python3
"""
创建包含UML图像的完整DOCX文档
使用PlantUML生成UML图并嵌入Word文档
"""

import os
import re
import tempfile
import subprocess
import hashlib
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys
import time

def setup_plantuml():
    """设置PlantUML环境"""
    jar_path = "/tmp/plantuml.jar"
    
    # 如果JAR文件不存在，下载它
    if not os.path.exists(jar_path):
        print("下载PlantUML JAR文件...")
        try:
            import urllib.request
            jar_url = "https://github.com/plantuml/plantuml/releases/download/v1.2023.12/plantuml-1.2023.12.jar"
            urllib.request.urlretrieve(jar_url, jar_path)
            print(f"下载完成: {jar_path}")
        except Exception as e:
            print(f"下载失败: {e}")
            # 尝试备用URL
            try:
                jar_url = "https://sourceforge.net/projects/plantuml/files/plantuml.jar/download"
                urllib.request.urlretrieve(jar_url, jar_path)
                print(f"从备用源下载完成: {jar_path}")
            except:
                print("无法下载PlantUML JAR文件")
                return None
    
    return jar_path

def generate_uml_image(plantuml_code, output_dir, index, jar_path):
    """生成UML图像"""
    # 创建临时文件
    temp_dir = tempfile.mkdtemp()
    temp_file = os.path.join(temp_dir, f"uml_{index}.puml")
    
    try:
        # 写入PlantUML代码
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write(plantuml_code)
        
        # 生成图像
        output_file = os.path.join(output_dir, f'uml_{index}.png')
        
        cmd = ['java', '-jar', jar_path, '-tpng', temp_file, '-o', output_dir]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            # 检查生成的文件
            expected_base = os.path.basename(temp_file).replace('.puml', '')
            for f in os.listdir(output_dir):
                if f.startswith(expected_base) and f.endswith('.png'):
                    generated_file = os.path.join(output_dir, f)
                    if os.path.exists(generated_file):
                        # 重命名为标准名称
                        os.rename(generated_file, output_file)
                        print(f"✅ 生成UML图 {index}: {output_file}")
                        return output_file
        else:
            print(f"⚠️ PlantUML错误 (图{index}): {result.stderr[:100]}")
            
    except subprocess.TimeoutExpired:
        print(f"⏱️  UML图 {index} 生成超时")
    except Exception as e:
        print(f"❌ UML图 {index} 生成失败: {e}")
    finally:
        # 清理临时文件
        try:
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
        except:
            pass
    
    return None

def extract_markdown_sections(md_content):
    """提取Markdown文档的各个部分和PlantUML代码块"""
    
    # 正则表达式匹配PlantUML代码块
    plantuml_pattern = r'```plantuml\s*(.*?)```'
    
    # 分割文档
    sections = []
    current_section = ""
    in_code_block = False
    code_block_type = ""
    code_content = []
    
    lines = md_content.split('\n')
    
    for line in lines:
        # 检查代码块开始
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_type = line.strip()[3:].strip()
                code_content = []
                
                # 保存当前段落
                if current_section.strip():
                    sections.append(('text', current_section))
                    current_section = ""
            else:
                in_code_block = False
                
                if code_block_type == 'plantuml':
                    # 保存PlantUML代码块
                    sections.append(('plantuml', '\n'.join(code_content)))
                else:
                    # 保存其他代码块为文本
                    sections.append(('text', f'```{code_block_type}\n' + '\n'.join(code_content) + '\n```'))
                
                continue
        
        if in_code_block:
            code_content.append(line)
        else:
            current_section += line + '\n'
    
    # 添加最后一段
    if current_section.strip():
        sections.append(('text', current_section))
    
    return sections

def create_docx_with_uml_images(md_file, docx_file):
    """创建包含UML图像的DOCX文档"""
    
    print(f"📖 读取Markdown文档: {md_file}")
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 设置PlantUML
    print("🔧 设置PlantUML环境...")
    jar_path = setup_plantuml()
    if not jar_path:
        print("❌ 无法设置PlantUML环境")
        return False
    
    # 创建临时目录存储UML图像
    temp_image_dir = tempfile.mkdtemp()
    print(f"🖼️ 图像临时目录: {temp_image_dir}")
    
    # 提取文档部分
    print("📋 解析Markdown文档...")
    sections = extract_markdown_sections(md_content)
    
    # 统计PlantUML代码块
    plantuml_blocks = [content for type_, content in sections if type_ == 'plantuml']
    print(f"📊 发现 {len(plantuml_blocks)} 个PlantUML代码块")
    
    # 生成UML图像
    print("🎨 生成UML图像...")
    image_files = []
    for i, code in enumerate(plantuml_blocks):
        print(f"  处理UML图 {i+1}/{len(plantuml_blocks)}...")
        image_file = generate_uml_image(code, temp_image_dir, i+1, jar_path)
        image_files.append(image_file)
    
    # 创建Word文档
    print("📄 创建Word文档...")
    doc = Document()
    
    # 设置文档样式
    styles = doc.styles
    style = styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 添加自定义样式
    if 'Heading 1' not in styles:
        heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 0)
    
    if 'Heading 2' not in styles:
        heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(30, 30, 30)
    
    if 'Heading 3' not in styles:
        heading3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.font.color.rgb = RGBColor(60, 60, 60)
    
    # 处理文档部分
    print("✍️ 写入文档内容...")
    uml_index = 0
    
    for section_type, content in sections:
        if section_type == 'text':
            # 处理文本内容
            lines = content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                if not line.strip():
                    # 空行
                    doc.add_paragraph()
                    continue
                
                # 处理标题
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('##### '):
                    doc.add_heading(line[6:], level=5)
                elif line.startswith('###### '):
                    doc.add_heading(line[7:], level=6)
                
                # 处理列表
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line.strip()):
                    doc.add_paragraph(line.strip(), style='List Number')
                
                # 处理表格行（简化）
                elif '|' in line and '---' not in line and line.count('|') >= 2:
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        # 简单处理，实际应该解析完整表格
                        para = doc.add_paragraph()
                        para.add_run(' | '.join(cells))
                
                # 普通段落
                elif line.strip() and not line.startswith('#'):
                    doc.add_paragraph(line.strip())
        
        elif section_type == 'plantuml':
            # 添加UML图像
            uml_index += 1
            
            if uml_index <= len(image_files) and image_files[uml_index-1]:
                image_path = image_files[uml_index-1]
                
                try:
                    # 添加图像标题
                    title_para = doc.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title_para.add_run(f'图 {uml_index}: UML图')
                    run.bold = True
                    run.font.size = Pt(11)
                    
                    # 添加图像
                    doc.add_picture(image_path, width=Inches(6.5))
                    
                    # 添加图像说明
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f'[UML图 {uml_index}]')
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
                    caption_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    doc.add_paragraph()  # 空行
                    
                    print(f"  已插入UML图 {uml_index}")
                    
                except Exception as e:
                    print(f"⚠️ 插入UML图 {uml_index} 失败: {e}")
                    # 添加占位符
                    doc.add_paragraph(f"[UML图 {uml_index} - 图像加载失败]")
                    doc.add_paragraph("PlantUML代码:")
                    doc.add_paragraph(content[:200] + '...' if len(content) > 200 else content)
            else:
                # 图像生成失败
                doc.add_paragraph(f"[UML图 {uml_index} - 生成失败]")
                doc.add_paragraph("PlantUML代码:")
                doc.add_paragraph(content[:200] + '...' if len(content) > 200 else content)
    
    # 保存文档
    print(f"💾 保存文档: {docx_file}")
    doc.save(docx_file)
    
    # 清理临时文件
    try:
        import shutil
        shutil.rmtree(temp_image_dir, ignore_errors=True)
        print("🧹 清理临时文件")
    except:
        pass
    
    return True

def main():
    print("=" * 60)
    print("孕期宝UML文档转换工具")
    print("使用PlantUML生成图像并嵌入Word文档")
    print("=" * 60)
    
    # 文件路径
    md_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝架构设计文档.md'
    docx_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝UML完整文档.docx'
    
    if not os.path.exists(md_file):
        print(f"❌ 错误: Markdown文件不存在: {md_file}")
        return
    
    print(f"📄 输入文件: {md_file}")
    print(f"📄 输出文件: {docx_file}")
    print(f"📊 文件大小: {os.path.getsize(md_file):,} 字节")
    
    start_time = time.time()
    
    try:
        success = create_docx_with_uml_images(md_file, docx_file)
        
        if success:
            end_time = time.time()
            elapsed = end_time - start_time
            
            if os.path.exists(docx_file):
                size = os.path.getsize(docx_file)
                print("=" * 60)
                print("✅ 转换成功!")
                print(f"📁 输出文件: {docx_file}")
                print(f"💾 文件大小: {size:,} 字节 ({size/1024:.1f} KB)")
                print(f"⏱️ 耗时: {elapsed:.1f} 秒")
                print("=" * 60)
                
                # 显示统计信息
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    uml_count = content.count('```plantuml')
                    print(f"📈 统计:")
                    print(f"   - UML图数量: {uml_count}")
                    print(f"   - 文档页数: 约 {size//5000 + 1} 页")
            else:
                print("❌ 错误: 输出文件未创建")
        else:
            print("❌ 转换失败")
            
    except Exception as e:
        print(f"❌ 转换过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()