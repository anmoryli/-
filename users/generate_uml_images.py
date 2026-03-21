#!/usr/bin/env python3
"""
生成PlantUML图像并创建DOCX文档
"""

import os
import re
import tempfile
import subprocess
import hashlib
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_plantuml_jar():
    """下载PlantUML JAR文件"""
    jar_url = "https://github.com/plantuml/plantuml/releases/download/v1.2023.12/plantuml-1.2023.12.jar"
    jar_path = "/tmp/plantuml.jar"
    
    if not os.path.exists(jar_path):
        print("下载PlantUML JAR文件...")
        try:
            import urllib.request
            urllib.request.urlretrieve(jar_url, jar_path)
            print(f"下载完成: {jar_path}")
        except Exception as e:
            print(f"下载失败: {e}")
            return None
    
    return jar_path

def generate_uml_image(plantuml_code, output_dir, index):
    """使用PlantUML JAR生成图像"""
    # 创建临时PlantUML文件
    temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.puml', encoding='utf-8')
    temp_file.write(plantuml_code)
    temp_file.close()
    
    # 生成图像
    jar_path = download_plantuml_jar()
    if not jar_path:
        return None
    
    output_file = os.path.join(output_dir, f'uml_{index}.png')
    
    try:
        cmd = ['java', '-jar', jar_path, '-tpng', temp_file.name, '-o', output_dir]
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            # 检查生成的文件
            expected_file = os.path.join(output_dir, os.path.basename(temp_file.name).replace('.puml', '.png'))
            if os.path.exists(expected_file):
                os.rename(expected_file, output_file)
                print(f"生成图像: {output_file}")
                return output_file
            else:
                # 尝试其他可能的文件名
                for f in os.listdir(output_dir):
                    if f.endswith('.png') and f.startswith(os.path.basename(temp_file.name).replace('.puml', '')):
                        os.rename(os.path.join(output_dir, f), output_file)
                        return output_file
        else:
            print(f"PlantUML错误: {result.stderr}")
    except Exception as e:
        print(f"生成图像失败: {e}")
    finally:
        # 清理临时文件
        try:
            os.unlink(temp_file.name)
        except:
            pass
    
    return None

def extract_plantuml_blocks(md_content):
    """从Markdown中提取PlantUML代码块"""
    pattern = r'```plantuml\s*(.*?)```'
    blocks = re.findall(pattern, md_content, re.DOTALL)
    return blocks

def create_docx_with_images(md_file, docx_file):
    """创建包含UML图像的DOCX文档"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建输出目录
    output_dir = tempfile.mkdtemp()
    print(f"临时目录: {output_dir}")
    
    # 提取PlantUML代码块
    plantuml_blocks = extract_plantuml_blocks(md_content)
    print(f"找到 {len(plantuml_blocks)} 个PlantUML代码块")
    
    # 生成图像
    image_files = []
    for i, block in enumerate(plantuml_blocks):
        print(f"生成UML图 {i+1}/{len(plantuml_blocks)}...")
        image_file = generate_uml_image(block, output_dir, i+1)
        if image_file:
            image_files.append(image_file)
        else:
            print(f"警告: 无法生成UML图 {i+1}")
            image_files.append(None)
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 处理Markdown内容
    lines = md_content.split('\n')
    in_code_block = False
    code_block_type = ''
    current_code = []
    uml_index = 0
    
    for line in lines:
        # 检查代码块开始
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_type = line.strip()[3:].strip()
                current_code = []
            else:
                in_code_block = False
                
                # 处理代码块
                if code_block_type == 'plantuml':
                    # 添加UML图像
                    if uml_index < len(image_files) and image_files[uml_index]:
                        try:
                            # 添加图像
                            doc.add_picture(image_files[uml_index], width=Inches(6))
                            
                            # 添加图像说明
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(f'[UML图 {uml_index+1}]')
                            run.italic = True
                            run.font.size = Pt(9)
                        except Exception as e:
                            print(f"添加图像失败: {e}")
                            doc.add_paragraph(f"[UML图 {uml_index+1} - 图像加载失败]")
                    else:
                        doc.add_paragraph(f"[UML图 {uml_index+1} - 生成失败]")
                        # 添加部分代码作为参考
                        code_preview = '\n'.join(current_code[:5])
                        if len(current_code) > 5:
                            code_preview += '\n...'
                        doc.add_paragraph(code_preview)
                    
                    uml_index += 1
                else:
                    # 普通代码块
                    doc.add_paragraph(f"[代码块 - {code_block_type}]")
                    code_text = doc.add_paragraph()
                    code_text.add_run('\n'.join(current_code))
                
                doc.add_paragraph()  # 空行
                continue
        
        if in_code_block:
            current_code.append(line)
            continue
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            doc.add_heading(line[7:], level=6)
        
        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            doc.add_paragraph(line.strip(), style='List Number')
        
        # 处理表格标题行（简化）
        elif '|' in line and '---' not in line and line.count('|') >= 2:
            # 简单表格处理
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                # 这里简化处理表格
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid'
                hdr_cells = table.rows[0].cells
                for i, cell in enumerate(cells):
                    hdr_cells[i].text = cell
        
        # 普通段落
        elif line.strip():
            doc.add_paragraph(line.strip())
        elif not line.strip():
            # 空行
            doc.add_paragraph()
    
    # 保存文档
    doc.save(docx_file)
    print(f"文档已保存: {docx_file}")
    
    # 清理临时文件
    try:
        import shutil
        shutil.rmtree(output_dir)
        print("清理临时文件")
    except:
        pass

def main():
    md_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝UML文档.md'
    docx_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝UML文档.docx'
    
    if not os.path.exists(md_file):
        print(f"错误: Markdown文件不存在: {md_file}")
        return
    
    print(f"开始转换: {md_file} -> {docx_file}")
    
    try:
        create_docx_with_images(md_file, docx_file)
        
        if os.path.exists(docx_file):
            print(f"转换完成! 文件大小: {os.path.getsize(docx_file)} 字节")
        else:
            print("错误: 输出文件未创建")
            
    except Exception as e:
        print(f"转换过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()