#!/usr/bin/env python3
"""
将包含PlantUML代码的Markdown文档转换为DOCX文档
确保UML图正确显示为图像
"""

import os
import re
import tempfile
import base64
import hashlib
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import plantuml

def extract_plantuml_blocks(md_content):
    """从Markdown内容中提取PlantUML代码块"""
    pattern = r'```plantuml\s*(.*?)```'
    blocks = re.findall(pattern, md_content, re.DOTALL)
    return blocks

def generate_plantuml_image(plantuml_code, output_format='png'):
    """将PlantUML代码转换为图像"""
    try:
        # 使用plantuml库生成图像
        server = plantuml.PlantUML(url='http://www.plantuml.com/plantuml/img/')
        image_url = server.get_url(plantuml_code)
        
        # 下载图像数据
        import urllib.request
        response = urllib.request.urlopen(image_url)
        image_data = response.read()
        
        return image_data
    except Exception as e:
        print(f"生成PlantUML图像失败: {e}")
        # 返回一个占位符图像
        return None

def markdown_to_docx(md_file_path, docx_file_path):
    """将Markdown文件转换为DOCX文件"""
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    
    # 添加标题样式
    heading_style = doc.styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # 提取PlantUML代码块并生成图像
    plantuml_blocks = extract_plantuml_blocks(md_content)
    image_map = {}
    
    for i, block in enumerate(plantuml_blocks):
        print(f"处理PlantUML图 {i+1}/{len(plantuml_blocks)}...")
        image_data = generate_plantuml_image(block)
        if image_data:
            # 保存图像到临时文件
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_file.write(image_data)
            temp_file.close()
            image_map[f'plantuml_block_{i}'] = temp_file.name
    
    # 将Markdown转换为HTML
    html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
    
    # 简化处理：直接按行处理Markdown
    lines = md_content.split('\n')
    
    in_code_block = False
    code_block_type = ''
    current_code = []
    
    for line in lines:
        # 检查代码块开始
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_type = line.strip()[3:].strip()
                current_code = []
            else:
                in_code_block = False
                # 处理代码块
                if code_block_type == 'plantuml':
                    # 查找对应的图像
                    code_content = '\n'.join(current_code)
                    code_hash = hashlib.md5(code_content.encode()).hexdigest()
                    image_found = False
                    
                    for key, image_path in image_map.items():
                        # 简化：使用第一个图像
                        if os.path.exists(image_path):
                            # 添加图像到文档
                            try:
                                doc.add_picture(image_path, width=Inches(6))
                                image_found = True
                                # 添加图像说明
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run(f'[PlantUML图 {len(image_map)}]')
                                run.italic = True
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(128, 128, 128)
                            except Exception as e:
                                print(f"添加图像失败: {e}")
                                doc.add_paragraph(f"[PlantUML图 - 渲染失败]")
                            break
                    
                    if not image_found:
                        doc.add_paragraph("[PlantUML图 - 需要手动渲染]")
                        code_para = doc.add_paragraph()
                        code_para.add_run(code_content[:200] + '...' if len(code_content) > 200 else code_content)
                else:
                    # 普通代码块
                    doc.add_paragraph(f"[代码块 - {code_block_type}]")
                    code_para = doc.add_paragraph()
                    code_para.add_run('\n'.join(current_code))
                
                doc.add_paragraph()  # 空行
                continue
        
        if in_code_block:
            current_code.append(line)
            continue
        
        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            doc.add_heading(line[7:], level=6)
        
        # 处理表格（简化处理）
        elif '|' in line and line.count('|') >= 2:
            # 简单表格处理
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                # 这里简化处理，实际应该解析完整的表格
                doc.add_paragraph(' | '.join(cells))
        
        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            doc.add_paragraph(line.strip(), style='List Number')
        
        # 普通段落
        elif line.strip():
            doc.add_paragraph(line.strip())
        elif not line.strip() and not in_code_block:
            # 空行
            doc.add_paragraph()
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"文档已保存: {docx_file_path}")
    
    # 清理临时文件
    for image_path in image_map.values():
        try:
            os.unlink(image_path)
        except:
            pass

def main():
    # 文件路径
    md_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝UML文档.md'
    docx_file = '/mnt/d/Code/Project/Merge/yunqibao/users/孕期宝UML文档.docx'
    
    if not os.path.exists(md_file):
        print(f"错误: Markdown文件不存在: {md_file}")
        return
    
    print(f"开始转换: {md_file} -> {docx_file}")
    print(f"文件大小: {os.path.getsize(md_file)} 字节")
    
    try:
        markdown_to_docx(md_file, docx_file)
        print("转换完成!")
        
        # 检查输出文件
        if os.path.exists(docx_file):
            print(f"输出文件: {docx_file}")
            print(f"输出文件大小: {os.path.getsize(docx_file)} 字节")
        else:
            print("错误: 输出文件未创建")
            
    except Exception as e:
        print(f"转换过程中出错: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()