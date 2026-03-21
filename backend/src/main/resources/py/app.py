# -*- coding: utf-8 -*-
# 文件名: app.py
import os
import time
import logging
import re
from typing import List, Optional
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, HTTPException, Query, Body
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from openai import OpenAI
from pymilvus import MilvusClient, DataType

from docx import Document as DocxDocument
from tqdm import tqdm
import uvicorn

# ==================== 日志 ====================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    handlers=[logging.FileHandler("rag.log", encoding="utf-8"), logging.StreamHandler()]
)
logger = logging.getLogger("RAG")

# ==================== 配置 ====================
# user_id = -1 表示全局百科；>0 表示对应用户私有数据
GLOBAL_USER_ID = -1


class Config:
    MILVUS_URI = "http://localhost:19530"
    MILVUS_TOKEN = "root:Milvus"
    COLLECTION_NAME = "yunji"
    RAG_COLLECTION_NAME = "yunji_rag"  # 用户+百科统一集合，按 user_id 区分
    VECTOR_DIM = 1024

    CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 100

    EMBEDDING_MODEL = "text-embedding-v3"
    LLM_MODEL = "qwen-plus"

    POLICY_DIR = Path("./knowledge")
    DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY")

    if not DASHSCOPE_API_KEY:
        raise RuntimeError("请设置环境变量 DASHSCOPE_API_KEY")

config = Config()

# ==================== FastAPI ====================
app = FastAPI(title="孕记 RAG 系统", version="2.3.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True,
                   allow_methods=["*"], allow_headers=["*"])

# ==================== 客户端 ====================
class Clients:
    def __init__(self):
        self.milvus = MilvusClient(uri=config.MILVUS_URI, token=config.MILVUS_TOKEN)
        self.embedder = OpenAI(api_key=config.DASHSCOPE_API_KEY,
                               base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
        self.llm = OpenAI(api_key=config.DASHSCOPE_API_KEY,
                          base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
        self._ensure_collection()

    def _ensure_collection(self):
        if not self.milvus.has_collection(config.COLLECTION_NAME):
            self.milvus.create_collection(
                collection_name=config.COLLECTION_NAME,
                dimension=config.VECTOR_DIM,
                metric_type="IP",
                consistency_level="Strong",
                auto_id=True,
                primary_field="id",
                vector_field="vector",
                description="政策原文块直接向量化"
            )
            logger.info(f"创建集合 {config.COLLECTION_NAME}")
        if not self.milvus.has_collection(config.RAG_COLLECTION_NAME):
            schema = self.milvus.create_schema(auto_id=True, enable_dynamic_field=False)
            schema.add_field(field_name="id", datatype=DataType.INT64, is_primary=True, auto_id=True)
            schema.add_field(field_name="vector", datatype=DataType.FLOAT_VECTOR, dim=config.VECTOR_DIM)
            schema.add_field(field_name="text", datatype=DataType.VARCHAR, max_length=65535)
            schema.add_field(field_name="user_id", datatype=DataType.INT64)
            schema.add_field(field_name="source", datatype=DataType.VARCHAR, max_length=256)
            schema.add_field(field_name="source_id", datatype=DataType.VARCHAR, max_length=256)
            schema.add_field(field_name="upload_time", datatype=DataType.VARCHAR, max_length=64)
            idx = self.milvus.prepare_index_params()
            idx.add_index(field_name="vector", index_type="AUTOINDEX", index_name="vector_idx", metric_type="IP", params={})
            self.milvus.create_collection(
                collection_name=config.RAG_COLLECTION_NAME,
                schema=schema,
                index_params=idx,
            )
            logger.info(f"创建集合 {config.RAG_COLLECTION_NAME}（含 user_id）")

clients = Clients()

# ==================== 文本清理（关键！）===================
def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'[-—=]+\s*\d+\s*[-—=]+', ' ', text)   # 去页眉 —1—
    text = re.sub(r'第\s*\d+\s*页|共\s*\d+\s*页', ' ', text)
    text = re.sub(r'\n\s*\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\r', '', text)
    text = text.strip()

    lines = [line.strip() for line in text.split('\n') if line.strip() and not re.fullmatch(r'\d+', line.strip())]
    return '\n'.join(lines)

# ==================== 分块 ====================
splitter = RecursiveCharacterTextSplitter(
    chunk_size=config.CHUNK_SIZE,
    chunk_overlap=config.CHUNK_OVERLAP,
    length_function=len,
    separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""]
)

def load_and_split(file_path: Path) -> List[Document]:
    try:
        if file_path.suffix.lower() == ".pdf":
            loader = PyPDFLoader(str(file_path))
            docs = loader.load()
        elif file_path.suffix.lower() == ".docx":
            doc = DocxDocument(str(file_path))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            docs = [Document(page_content=text)]
        elif file_path.suffix.lower() == ".md":
            text = file_path.read_text(encoding="utf-8")
            docs = [Document(page_content=text)]
        else:
            return []

        chunks = splitter.split_documents(docs)
        logger.info(f"{file_path.name} 分块完成 → {len(chunks)} 块")
        return chunks
    except Exception as e:
        logger.error(f"加载失败 {file_path}: {e}")
        return []

# ==================== 一键构建向量库，是把所以相关的资料全部上传到服务器的/knowledge文件夹下，可以是pdf或者docx/.md ====================
@app.post("/api/v1/upload")
async def upload_and_build():
    start = time.time()
    entities = []
    total_chunks = 0

    if not config.POLICY_DIR.exists():
        raise HTTPException(400, "knowledge/ 目录不存在")

    files = [f for f in config.POLICY_DIR.iterdir() if f.suffix.lower() in {".pdf", ".docx"}]
    if not files:
        raise HTTPException(400, "knowledge/ 目录没有 PDF/DOCX 文件")

    for file in tqdm(files, desc="构建向量库"):
        chunks = load_and_split(file)
        for i, chunk in enumerate(chunks):
            raw = chunk.page_content
            cleaned = clean_text(raw)
            if not cleaned:
                continue

            vector = clients.embedder.embeddings.create(
                model=config.EMBEDDING_MODEL,
                input=cleaned,
                dimensions=config.VECTOR_DIM
            ).data[0].embedding

            entities.append({
                "vector": vector,
                "text": cleaned,
                "chunk_id": i,
                "upload_time": datetime.now().isoformat(),
                "source": file.name
            })
            total_chunks += 1

    if entities:
        res = clients.milvus.insert(config.COLLECTION_NAME, entities)
        logger.info(f"成功插入 {res.insert_count} 条向量")

    return {
        "code": 200,
        "message": "向量库构建完成",
        "data": {
            "files": len(files),
            "chunks": total_chunks,
            "vectors": len(entities),
            "latency_ms": round((time.time() - start) * 1000, 2)
        }
    }

# ==================== 问题重写 ====================
async def rewrite_query(q: str) -> str:
    prompt = f"将下面用户问题重写为适合向量检索的政策问题，只输出一句话：\n{q}"
    try:
        r = clients.llm.chat.completions.create(
            model=config.LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=60
        )
        return r.choices[0].message.content.strip()
    except:
        return q

# ==================== 单条嵌入（Java 调用：memo/消息/PDF/图片描述）===================
class EmbedBody(BaseModel):
    user_id: int
    text: str
    source: str = "memo"   # memo | message | pdf | image_desc
    source_id: Optional[str] = None


@app.post("/api/v1/embed")
async def embed_one(body: EmbedBody = Body(...)):
    if not body.text or not body.text.strip():
        return {"code": 400, "message": "text 不能为空", "data": None}
    text = body.text.strip()[:10000]
    try:
        vector = clients.embedder.embeddings.create(
            model=config.EMBEDDING_MODEL,
            input=text,
            dimensions=config.VECTOR_DIM
        ).data[0].embedding
        entity = {
            "vector": vector,
            "text": text,
            "user_id": body.user_id,
            "source": body.source,
            "source_id": body.source_id or "",
            "upload_time": datetime.now().isoformat(),
        }
        clients.milvus.insert(config.RAG_COLLECTION_NAME, [entity])
        logger.info(f"嵌入成功 user_id={body.user_id} source={body.source} source_id={body.source_id}")
        return {"code": 200, "message": "ok", "data": {"inserted": 1}}
    except Exception as e:
        logger.error(f"嵌入失败: {e}")
        raise HTTPException(500, str(e))


# ==================== 按 source + source_id 删除向量（多源异构数据同步）===================
@app.delete("/api/v1/delete")
async def delete_vectors(
        source: str = Query(..., description="memo|message|pdf|image_desc"),
        source_id: str = Query(..., description="业务表主键"),
):
    """按 source + source_id 删除向量库中的对应条目，用于记录删除时同步清理"""
    try:
        # Milvus 过滤表达式：source 和 source_id 为 VARCHAR 类型，需用引号
        filter_expr = f'source == "{source}" and source_id == "{source_id}"'
        clients.milvus.delete(collection_name=config.RAG_COLLECTION_NAME, filter=filter_expr)
        logger.info(f"删除向量成功 source={source} source_id={source_id}")
        return {"code": 200, "message": "deleted", "data": None}
    except Exception as e:
        logger.error(f"删除向量失败 source={source} source_id={source_id}: {e}")
        raise HTTPException(500, str(e))


# ==================== 百科嵌入（knowledge 目录 → 全局区 user_id=-1）===================
@app.post("/api/v1/embed-knowledge")
async def embed_knowledge():
    start = time.time()
    entities = []
    if not config.POLICY_DIR.exists():
        raise HTTPException(400, "knowledge/ 目录不存在")
    files = [f for f in config.POLICY_DIR.iterdir() if f.suffix.lower() in {".pdf", ".docx", ".md"}]
    if not files:
        raise HTTPException(400, "knowledge/ 目录没有 PDF/DOCX/MD 文件")
    for file in tqdm(files, desc="百科嵌入"):
        chunks = load_and_split(file)
        for i, chunk in enumerate(chunks):
            raw = chunk.page_content
            cleaned = clean_text(raw)
            if not cleaned:
                continue
            vector = clients.embedder.embeddings.create(
                model=config.EMBEDDING_MODEL,
                input=cleaned,
                dimensions=config.VECTOR_DIM
            ).data[0].embedding
            entities.append({
                "vector": vector,
                "text": cleaned,
                "user_id": GLOBAL_USER_ID,
                "source": file.name,
                "source_id": "",
                "upload_time": datetime.now().isoformat(),
            })
    if entities:
        clients.milvus.insert(config.RAG_COLLECTION_NAME, entities)
    latency = round((time.time() - start) * 1000, 2)
    return {"code": 200, "message": "百科嵌入完成", "data": {"chunks": len(entities), "latency_ms": latency}}


def _search_rag_collection(q_vec: list, top_k: int, user_id_filter: Optional[int] = None) -> list:
    """在 yunji_rag 中检索。user_id_filter=None 只查百科；否则先查用户区再查百科，合并后按分数排序取 top_k"""
    all_hits = []
    try:
        output_f = ["text", "user_id", "source", "source_id", "upload_time"]
        limit_per = min(top_k, 20)
        if user_id_filter is not None and user_id_filter > 0:
            for uid in [user_id_filter, GLOBAL_USER_ID]:
                search_res = clients.milvus.search(
                    collection_name=config.RAG_COLLECTION_NAME,
                    data=[q_vec],
                    limit=limit_per,
                    output_fields=output_f,
                    search_params={"metric_type": "IP", "params": {}},
                    filter=f"user_id == {uid}",
                )
                for result_set in search_res:
                    for hit in result_set:
                        entity = hit.entity
                        full_text = (entity.get("text") or "").strip()
                        if not full_text:
                            continue
                        display = clean_text(full_text)
                        if len(display) > 1500:
                            display = display[:1500] + "..."
                        meta = {k: entity.get(k) for k in output_f}
                        all_hits.append({"similarity_score": round(hit.distance, 4), "context": display, "metadata": meta})
        else:
            search_res = clients.milvus.search(
                collection_name=config.RAG_COLLECTION_NAME,
                data=[q_vec],
                limit=limit_per,
                output_fields=output_f,
                search_params={"metric_type": "IP", "params": {}},
                filter=f"user_id == {GLOBAL_USER_ID}",
            )
            for result_set in search_res:
                for hit in result_set:
                    entity = hit.entity
                    full_text = (entity.get("text") or "").strip()
                    if not full_text:
                        continue
                    display = clean_text(full_text)
                    if len(display) > 1500:
                        display = display[:1500] + "..."
                    meta = {k: entity.get(k) for k in output_f}
                    all_hits.append({"similarity_score": round(hit.distance, 4), "context": display, "metadata": meta})
        all_hits.sort(key=lambda x: x["similarity_score"], reverse=True)
        hits = []
        for i, h in enumerate(all_hits[:top_k]):
            hits.append({"rank": i + 1, "similarity_score": h["similarity_score"], "context": h["context"], "metadata": h["metadata"]})
        return hits
    except Exception as e:
        logger.warning(f"RAG 检索异常: {e}")
        return []


# ==================== 仅检索百科（全局区）===================
@app.get("/api/v1/query-knowledge")
async def query_knowledge(
        question: str = Query(..., description="用户问题"),
        top_k: int = Query(10, ge=1, le=20),
):
    start = time.time()
    optimized = await rewrite_query(question)
    q_vec = clients.embedder.embeddings.create(
        model=config.EMBEDDING_MODEL,
        input=optimized,
        dimensions=config.VECTOR_DIM
    ).data[0].embedding
    hits = _search_rag_collection(q_vec, top_k, user_id_filter=None)
    latency = round((time.time() - start) * 1000, 2)
    return {"code": 200, "message": "success", "data": {"original_question": question, "optimized_query": optimized, "latency_ms": latency, "hit_count": len(hits), "hits": hits}}


# ==================== 查询接口（支持 user_id、include_global，Java 直接拿 hits）===================
@app.get("/api/v1/query")
async def query(
        question: str = Query(..., description="用户问题"),
        top_k: int = Query(10, ge=1, le=20, description="返回条数，默认10，Java可传3"),
        user_id: Optional[int] = Query(None, description="当前用户ID，不传则仅查百科"),
        include_global: bool = Query(True, description="是否同时检索百科（user_id=-1）"),
):
    start = time.time()
    optimized = await rewrite_query(question)
    q_vec = clients.embedder.embeddings.create(
        model=config.EMBEDDING_MODEL,
        input=optimized,
        dimensions=config.VECTOR_DIM
    ).data[0].embedding

    use_rag = clients.milvus.has_collection(config.RAG_COLLECTION_NAME)
    if use_rag:
        effective_user_id = user_id if (user_id is not None and user_id > 0 and include_global) else (user_id if (user_id is not None and user_id > 0) else None)
        if effective_user_id is None and not include_global:
            hits = []
        else:
            hits = _search_rag_collection(q_vec, top_k, user_id_filter=effective_user_id if (user_id is not None and user_id > 0) else None)
    else:
        search_res = clients.milvus.search(
            collection_name=config.COLLECTION_NAME,
            data=[q_vec],
            limit=top_k,
            output_fields=["text", "chunk_id", "upload_time", "source"],
            search_params={"metric_type": "IP", "params": {}},
        )
        hits = []
        for result_set in search_res:
            for hit in result_set:
                entity = hit.entity
                full_text = entity.get("text") or ""
                display = clean_text(full_text)
                if len(display) > 1500:
                    display = display[:1500] + "..."
                hits.append({"rank": len(hits) + 1, "similarity_score": round(hit.distance, 4), "context": display, "metadata": {"chunk_id": entity.get("chunk_id"), "upload_time": entity.get("upload_time"), "source": entity.get("source")}})

    latency = round((time.time() - start) * 1000, 2)
    logger.info(f"查询「{question}」user_id={user_id} include_global={include_global} 耗时 {latency}ms，命中 {len(hits)} 条")
    return {"code": 200, "message": "success", "data": {"original_question": question, "optimized_query": optimized, "latency_ms": latency, "hit_count": len(hits), "hits": hits}}

# ==================== 健康检查 ====================
@app.get("/health")
async def health():
    return {"status": "ok", "time": datetime.now().isoformat()}

# ==================== 启动 ====================
if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8004, log_level="info")